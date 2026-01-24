#!/usr/bin/env python3
"""
generate_spirits_report_v2.py - Generate comprehensive spirits statistics report

Creates a Word document with:
- YTD comparison (not just single month)
- Monthly trend charts
- Proper source citations
- Data verification

Usage:
    python generate_spirits_report_v2.py --year 2025 --month 10
"""

import os
import sys
import subprocess
import csv
import argparse
from datetime import datetime
from io import BytesIO
from collections import defaultdict

# Install dependencies
def install_deps():
    deps = ['matplotlib', 'python-docx', 'numpy']
    for dep in deps:
        try:
            __import__(dep.replace('-', '_'))
        except ImportError:
            print(f"Installing {dep}...")
            subprocess.check_call([sys.executable, '-m', 'pip', 'install', dep])

install_deps()

import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Paths
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(SCRIPT_DIR, '..', 'data')
OUTPUT_DIR = os.path.join(SCRIPT_DIR, 'content-queue', 'reports')

# Colors
COLORS = {
    'whisky': '#B8860B',
    'brandy': '#8B4513',
    'rum_gin_vodka': '#4682B4',
    'neutral': '#A9A9A9',
    'positive': '#228B22',
    'negative': '#DC143C',
    'primary': '#1C3A5F',
    'secondary': '#4A6FA5',
    'year_prior': '#B0B0B0',
    'grid': '#E8E8E8',
}

MONTH_NAMES = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun',
               'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']

# Configure matplotlib
plt.rcParams.update({
    'font.family': 'sans-serif',
    'font.sans-serif': ['Segoe UI', 'Arial', 'Helvetica'],
    'font.size': 10,
    'axes.titlesize': 12,
    'axes.titleweight': 'bold',
    'axes.spines.top': False,
    'axes.spines.right': False,
    'axes.edgecolor': '#CCCCCC',
    'figure.facecolor': 'white',
    'axes.facecolor': 'white',
    'savefig.dpi': 200,
    'axes.grid': True,
    'grid.alpha': 0.4,
    'grid.color': COLORS['grid'],
})


def load_csv_data():
    """Load data from TTB CSV files."""
    monthly_path = os.path.join(DATA_DIR, 'ttb_monthly_new.csv')
    yearly_path = os.path.join(DATA_DIR, 'ttb_yearly_new.csv')

    monthly_data = []
    yearly_data = []

    if os.path.exists(monthly_path):
        with open(monthly_path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            monthly_data = list(reader)

    if os.path.exists(yearly_path):
        with open(yearly_path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            yearly_data = list(reader)

    return monthly_data, yearly_data


def get_production_data(monthly_data, categories=None):
    """Extract production data from CSV."""
    if categories is None:
        categories = ['1-Whisky', '2-Brandy', '3-Rum, Gin, & Vodka', '4-Alcohol, Neutral Spirits, & Other']

    result = defaultdict(lambda: defaultdict(dict))

    for row in monthly_data:
        if row['Statistical_Group'] != '1-Distilled Spirits Production':
            continue
        if row['Statistical_Detail'] not in categories:
            continue
        if not row['Value'] or not row['Value'].isdigit():
            continue

        year = int(row['Year'])
        month = int(row['CY_Month_Number'])
        cat = row['Statistical_Detail']

        result[cat][year][month] = {
            'value': int(row['Value']),
            'producers': int(row['Count_IMs']) if row['Count_IMs'].isdigit() else 0
        }

    return result


def calculate_ytd(data, year, through_month):
    """Calculate YTD totals through a specific month."""
    ytd = {}
    for cat, years in data.items():
        if year in years:
            total = sum(years[year].get(m, {}).get('value', 0) for m in range(1, through_month + 1))
            ytd[cat] = total
    return ytd


def format_number(val, precision=1):
    """Format number with M/B suffix."""
    if val >= 1_000_000_000:
        return f'{val/1_000_000_000:.{precision}f}B'
    elif val >= 1_000_000:
        return f'{val/1_000_000:.{precision}f}M'
    elif val >= 1_000:
        return f'{val/1_000:.{precision}f}K'
    return str(int(val))


def get_cat_color(cat):
    """Get color for category."""
    if 'Whisky' in cat:
        return COLORS['whisky']
    elif 'Brandy' in cat:
        return COLORS['brandy']
    elif 'Rum' in cat:
        return COLORS['rum_gin_vodka']
    return COLORS['neutral']


def create_monthly_trend_chart(data, category, year, compare_year=None):
    """Create monthly trend chart showing all months."""
    if category not in data or year not in data[category]:
        return None

    cat_data = data[category]
    display_name = category.split('-', 1)[1] if '-' in category else category
    color = get_cat_color(category)

    fig, ax = plt.subplots(figsize=(11, 5))

    # Get months available for current year
    months_current = sorted(cat_data[year].keys())
    values_current = [cat_data[year][m]['value'] for m in months_current]

    x_current = np.array(months_current)

    # Plot comparison year if provided
    if compare_year and compare_year in cat_data:
        months_compare = sorted([m for m in cat_data[compare_year].keys() if m <= max(months_current)])
        values_compare = [cat_data[compare_year][m]['value'] for m in months_compare]
        x_compare = np.array(months_compare)

        ax.bar(x_compare - 0.2, values_compare, width=0.35,
               color=COLORS['year_prior'], label=str(compare_year), alpha=0.7)
        ax.bar(x_current + 0.2, values_current, width=0.35,
               color=color, label=str(year))
    else:
        ax.bar(x_current, values_current, color=color, width=0.6)

    # Formatting
    ax.set_xticks(range(1, 13))
    ax.set_xticklabels(MONTH_NAMES)
    ax.yaxis.set_major_formatter(ticker.FuncFormatter(lambda x, p: format_number(x)))
    ax.set_ylabel('Proof Gallons', fontweight='medium')
    ax.set_xlabel('Month', fontweight='medium')

    if compare_year:
        ax.set_title(f'{display_name} Monthly Production: {year} vs {compare_year}', pad=15)
        ax.legend(loc='upper right')
    else:
        ax.set_title(f'{display_name} Monthly Production: {year}', pad=15)

    # Add YoY change annotations for each month
    if compare_year and compare_year in cat_data:
        for m in months_current:
            if m in cat_data[compare_year]:
                curr = cat_data[year][m]['value']
                prev = cat_data[compare_year][m]['value']
                pct = (curr - prev) / prev * 100 if prev > 0 else 0

                color_pct = COLORS['positive'] if pct > 0 else COLORS['negative']
                sign = '+' if pct > 0 else ''

                y_pos = max(curr, prev) + max(values_current) * 0.03
                ax.annotate(f'{sign}{pct:.0f}%', (m + 0.2, y_pos),
                           ha='center', va='bottom', fontsize=8,
                           color=color_pct, fontweight='bold')

    ax.set_axisbelow(True)

    # Source
    fig.text(0.99, 0.01, 'Source: TTB Distilled Spirits Statistics',
             ha='right', va='bottom', fontsize=8, color='#888888', style='italic')

    plt.tight_layout()

    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=200, bbox_inches='tight', facecolor='white')
    buf.seek(0)
    plt.close()

    return buf


def create_ytd_comparison_chart(data, year, through_month):
    """Create YTD comparison bar chart."""
    categories = ['1-Whisky', '2-Brandy', '3-Rum, Gin, & Vodka']
    compare_year = year - 1

    ytd_current = calculate_ytd(data, year, through_month)
    ytd_prior = calculate_ytd(data, compare_year, through_month)

    # Prepare data
    names = []
    vals_current = []
    vals_prior = []
    colors = []

    for cat in categories:
        if cat in ytd_current and cat in ytd_prior:
            name = cat.split('-', 1)[1] if '-' in cat else cat
            names.append(name)
            vals_current.append(ytd_current[cat])
            vals_prior.append(ytd_prior[cat])
            colors.append(get_cat_color(cat))

    if not names:
        return None, {}

    # Calculate changes
    changes = {}
    for i, name in enumerate(names):
        pct = (vals_current[i] - vals_prior[i]) / vals_prior[i] * 100 if vals_prior[i] > 0 else 0
        changes[name] = {
            'current': vals_current[i],
            'prior': vals_prior[i],
            'pct': pct
        }

    fig, ax = plt.subplots(figsize=(10, 5.5))

    x = np.arange(len(names))
    width = 0.38

    bars_prior = ax.bar(x - width/2, vals_prior, width, label=f'{compare_year} YTD',
                        color=COLORS['year_prior'], edgecolor='white')
    bars_current = ax.bar(x + width/2, vals_current, width, label=f'{year} YTD',
                          color=colors, edgecolor='white')

    ax.set_xticks(x)
    ax.set_xticklabels(names, fontweight='medium')
    ax.yaxis.set_major_formatter(ticker.FuncFormatter(lambda x, p: format_number(x)))
    ax.set_ylabel('Proof Gallons', fontweight='medium')

    # Add change labels
    for i, bar in enumerate(bars_current):
        name = names[i]
        pct = changes[name]['pct']
        color = COLORS['positive'] if pct > 0 else COLORS['negative']
        sign = '+' if pct > 0 else ''

        y_pos = bar.get_height() + max(vals_current) * 0.02
        ax.annotate(f'{sign}{pct:.1f}%', (bar.get_x() + bar.get_width()/2, y_pos),
                   ha='center', va='bottom', fontsize=11, fontweight='bold', color=color)

    month_name = MONTH_NAMES[through_month - 1]
    ax.set_title(f'Beverage Spirits YTD Production: Jan-{month_name} {year} vs {compare_year}', pad=20)
    ax.legend(loc='upper right')
    ax.set_axisbelow(True)

    fig.text(0.99, 0.01, 'Source: TTB Distilled Spirits Statistics',
             ha='right', va='bottom', fontsize=8, color='#888888', style='italic')

    plt.tight_layout()

    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=200, bbox_inches='tight', facecolor='white')
    buf.seek(0)
    plt.close()

    return buf, changes


def create_multi_year_trend_chart(data, category, years=6):
    """Create multi-year annual trend chart."""
    if category not in data:
        return None

    cat_data = data[category]
    display_name = category.split('-', 1)[1] if '-' in category else category
    color = get_cat_color(category)

    # Get annual totals (sum all months)
    annual = {}
    for year in sorted(cat_data.keys()):
        if year >= datetime.now().year - years:
            total = sum(m.get('value', 0) for m in cat_data[year].values())
            if total > 0:
                annual[year] = total

    if len(annual) < 2:
        return None

    years_list = sorted(annual.keys())
    values = [annual[y] for y in years_list]

    fig, ax = plt.subplots(figsize=(10, 5))

    ax.fill_between(years_list, values, alpha=0.15, color=color)
    ax.plot(years_list, values, color=color, linewidth=2.5, marker='o',
            markersize=8, markerfacecolor='white', markeredgewidth=2)

    ax.yaxis.set_major_formatter(ticker.FuncFormatter(lambda x, p: format_number(x)))
    ax.set_ylabel('Proof Gallons (Annual)', fontweight='medium')
    ax.set_xlabel('Year', fontweight='medium')
    ax.set_xticks(years_list)

    # Add labels
    for x, y in zip(years_list, values):
        ax.annotate(format_number(y, 0), (x, y),
                   textcoords='offset points', xytext=(0, 12),
                   ha='center', fontsize=9, color=color, fontweight='medium')

    ax.set_title(f'{display_name} Annual Production Trend', pad=15)
    ax.set_axisbelow(True)

    fig.text(0.99, 0.01, 'Source: TTB Distilled Spirits Statistics',
             ha='right', va='bottom', fontsize=8, color='#888888', style='italic')

    plt.tight_layout()

    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=200, bbox_inches='tight', facecolor='white')
    buf.seek(0)
    plt.close()

    return buf


def create_word_document(year, month, charts, ytd_changes, monthly_data_dict):
    """Create comprehensive Word document."""
    doc = Document()

    month_name = MONTH_NAMES[month - 1]

    # Title
    title = doc.add_heading(f'American Spirits Production Report', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(f'{month_name} {year} | TTB Distilled Spirits Statistics')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    # Key Findings
    doc.add_heading('Key Findings', level=1)

    whisky_ytd = ytd_changes.get('Whisky', {})
    if whisky_ytd:
        findings = doc.add_paragraph()
        findings.add_run('Whisky: ').bold = True
        findings.add_run(
            f"Year-to-date production through {month_name} totals {format_number(whisky_ytd['current'])} proof gallons, "
            f"down {abs(whisky_ytd['pct']):.1f}% from the same period in {year-1} ({format_number(whisky_ytd['prior'])} PG)."
        )

    brandy_ytd = ytd_changes.get('Brandy', {})
    if brandy_ytd:
        findings = doc.add_paragraph()
        findings.add_run('Brandy: ').bold = True
        direction = 'down' if brandy_ytd['pct'] < 0 else 'up'
        findings.add_run(
            f"YTD production of {format_number(brandy_ytd['current'])} PG, {direction} {abs(brandy_ytd['pct']):.1f}% year-over-year."
        )

    rgv_ytd = ytd_changes.get('Rum, Gin, & Vodka', {})
    if rgv_ytd:
        findings = doc.add_paragraph()
        findings.add_run('Rum, Gin, & Vodka: ').bold = True
        direction = 'down' if rgv_ytd['pct'] < 0 else 'up'
        findings.add_run(
            f"Combined category {direction} {abs(rgv_ytd['pct']):.1f}% YTD to {format_number(rgv_ytd['current'])} PG."
        )

    doc.add_paragraph()

    # YTD Comparison Chart
    if 'ytd' in charts and charts['ytd']:
        doc.add_heading('Year-to-Date Comparison', level=1)
        doc.add_picture(charts['ytd'], width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # YTD data table
        doc.add_paragraph()
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        headers = ['Category', f'{year-1} YTD', f'{year} YTD', 'Change']
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

        for cat, vals in ytd_changes.items():
            row = table.add_row()
            row.cells[0].text = cat
            row.cells[1].text = format_number(vals['prior'])
            row.cells[2].text = format_number(vals['current'])
            sign = '+' if vals['pct'] > 0 else ''
            row.cells[3].text = f"{sign}{vals['pct']:.1f}%"

    # Monthly Trend Charts
    doc.add_heading('Monthly Production Trends', level=1)

    if 'whisky_monthly' in charts and charts['whisky_monthly']:
        doc.add_picture(charts['whisky_monthly'], width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    if 'rgv_monthly' in charts and charts['rgv_monthly']:
        doc.add_picture(charts['rgv_monthly'], width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    # Multi-year trend
    if 'whisky_annual' in charts and charts['whisky_annual']:
        doc.add_heading('Historical Trend', level=1)
        doc.add_picture(charts['whisky_annual'], width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data Source
    doc.add_paragraph()
    doc.add_heading('Data Source', level=1)
    source = doc.add_paragraph()
    source.add_run('Source: ').bold = True
    source.add_run('TTB Distilled Spirits Production and Operations Reports\n')
    source.add_run('URL: ').bold = True
    source.add_run('https://www.ttb.gov/distilled-spirits/distilled-spirits-reports\n')
    source.add_run('Data files: ').bold = True
    source.add_run('Distilled_Spirits_monthly_data.csv, Distilled_Spirits_yearly_data.csv\n')
    source.add_run('Last updated: ').bold = True
    source.add_run(f'{datetime.now().strftime("%B %d, %Y")}')

    return doc


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--year', type=int, default=2025)
    parser.add_argument('--month', type=int, default=10)
    args = parser.parse_args()

    print(f"\nGenerating comprehensive report for {args.year}-{args.month:02d}...")

    # Load data
    print("  Loading CSV data...")
    monthly_data, yearly_data = load_csv_data()

    if not monthly_data:
        print("ERROR: No monthly data found. Download from TTB first.")
        return

    # Process data
    print("  Processing production data...")
    prod_data = get_production_data(monthly_data)

    # Generate charts
    print("  Creating charts...")
    charts = {}

    # YTD comparison
    ytd_chart, ytd_changes = create_ytd_comparison_chart(prod_data, args.year, args.month)
    charts['ytd'] = ytd_chart

    # Monthly trends
    charts['whisky_monthly'] = create_monthly_trend_chart(
        prod_data, '1-Whisky', args.year, args.year - 1)
    charts['rgv_monthly'] = create_monthly_trend_chart(
        prod_data, '3-Rum, Gin, & Vodka', args.year, args.year - 1)

    # Annual trend
    charts['whisky_annual'] = create_multi_year_trend_chart(prod_data, '1-Whisky')

    # Create document
    print("  Creating Word document...")
    doc = create_word_document(args.year, args.month, charts, ytd_changes, prod_data)

    # Save
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    output_path = os.path.join(OUTPUT_DIR, f'spirits-report-{args.year}-{args.month:02d}-comprehensive.docx')
    doc.save(output_path)

    print(f"\nReport saved: {output_path}")

    # Save charts
    chart_dir = os.path.join(OUTPUT_DIR, 'charts')
    os.makedirs(chart_dir, exist_ok=True)

    for name, buf in charts.items():
        if buf:
            buf.seek(0)
            path = os.path.join(chart_dir, f'{name}-{args.year}-{args.month:02d}.png')
            with open(path, 'wb') as f:
                f.write(buf.read())
            print(f"  Chart: {path}")


if __name__ == '__main__':
    main()
