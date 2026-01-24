#!/usr/bin/env python3
"""
generate_spirits_report.py - Generate professional spirits statistics report with charts

Creates a Word document with charts and analysis commentary from TTB data.

Usage:
    python generate_spirits_report.py                    # Latest month
    python generate_spirits_report.py --year 2025 --month 10
    python generate_spirits_report.py --year 2024       # Annual report
"""

import os
import sys
import subprocess
import json
import argparse
from datetime import datetime
from io import BytesIO

# Install dependencies if needed
def install_deps():
    deps = ['matplotlib', 'python-docx']
    for dep in deps:
        try:
            __import__(dep.replace('-', '_').split('[')[0])
        except ImportError:
            print(f"Installing {dep}...")
            subprocess.check_call([sys.executable, '-m', 'pip', 'install', dep])

install_deps()

import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
from matplotlib.patches import Patch
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Paths
SCRIPT_DIR = os.path.dirname(__file__)
OUTPUT_DIR = os.path.join(SCRIPT_DIR, 'content-queue', 'reports')
WORKER_DIR = os.path.join(SCRIPT_DIR, '..', 'worker')

# Professional color palette
COLORS = {
    'whisky': '#B8860B',        # Dark goldenrod
    'brandy': '#8B4513',        # Saddle brown
    'rum_gin_vodka': '#4682B4', # Steel blue
    'neutral_spirits': '#A9A9A9', # Dark gray
    'positive': '#228B22',      # Forest green
    'negative': '#DC143C',      # Crimson
    'primary': '#1C3A5F',       # Navy blue
    'secondary': '#4A6FA5',     # Medium blue
    'text': '#2F2F2F',          # Dark gray text
    'grid': '#E5E5E5',          # Light gray grid
}

# Configure matplotlib for professional output
plt.rcParams.update({
    'font.family': 'sans-serif',
    'font.sans-serif': ['Segoe UI', 'Arial', 'Helvetica', 'DejaVu Sans'],
    'font.size': 11,
    'axes.titlesize': 13,
    'axes.titleweight': 'bold',
    'axes.labelsize': 10,
    'axes.labelweight': 'medium',
    'axes.spines.top': False,
    'axes.spines.right': False,
    'axes.edgecolor': '#CCCCCC',
    'axes.linewidth': 0.8,
    'xtick.labelsize': 9,
    'ytick.labelsize': 9,
    'legend.fontsize': 9,
    'figure.facecolor': 'white',
    'axes.facecolor': 'white',
    'savefig.facecolor': 'white',
    'savefig.dpi': 200,
    'axes.grid': True,
    'grid.alpha': 0.3,
    'grid.color': COLORS['grid'],
})


def run_d1_query(sql):
    """Execute D1 query via wrangler and return results."""
    sql_escaped = sql.replace('"', '\\"').replace('\n', ' ')
    cmd = f'npx wrangler d1 execute bevalc-colas --remote --json --command="{sql_escaped}"'
    result = subprocess.run(cmd, capture_output=True, text=True, cwd=WORKER_DIR, shell=True,
                           encoding='utf-8', errors='replace')

    if result.returncode == 0 and result.stdout:
        output = result.stdout
        start = output.find('[')
        if start >= 0:
            try:
                data = json.loads(output[start:])
                if data and isinstance(data, list) and len(data) > 0:
                    if 'results' in data[0]:
                        return data[0]['results']
            except json.JSONDecodeError:
                pass
    return []


def format_number(val, precision=1):
    """Format large numbers with M/B suffix."""
    if val >= 1_000_000_000:
        return f'{val/1_000_000_000:.{precision}f}B'
    elif val >= 1_000_000:
        return f'{val/1_000_000:.{precision}f}M'
    elif val >= 1_000:
        return f'{val/1_000:.{precision}f}K'
    return str(int(val))


def get_category_color(cat):
    """Get color for a category."""
    if 'Whisky' in cat:
        return COLORS['whisky']
    elif 'Brandy' in cat:
        return COLORS['brandy']
    elif 'Rum' in cat or 'Gin' in cat or 'Vodka' in cat:
        return COLORS['rum_gin_vodka']
    else:
        return COLORS['neutral_spirits']


def get_monthly_data(year, month):
    """Fetch all monthly production data."""
    sql = f"""
    SELECT statistical_detail, value, count_ims
    FROM ttb_spirits_stats
    WHERE year = {year} AND month = {month}
    AND statistical_group LIKE '1-Distilled Spirits Production%'
    AND statistical_detail IN ('1-Whisky', '2-Brandy', '3-Rum, Gin, & Vodka', '4-Alcohol, Neutral Spirits, & Other')
    ORDER BY value DESC
    """
    return run_d1_query(sql)


def get_yoy_data(year, month):
    """Fetch year-over-year comparison data."""
    sql = f"""
    SELECT year, statistical_detail, value, count_ims
    FROM ttb_spirits_stats
    WHERE month = {month}
    AND year IN ({year}, {year-1})
    AND statistical_group LIKE '1-Distilled Spirits Production%'
    AND statistical_detail IN ('1-Whisky', '2-Brandy', '3-Rum, Gin, & Vodka')
    ORDER BY statistical_detail, year
    """
    return run_d1_query(sql)


def get_yearly_trend(category, years=6):
    """Fetch multi-year trend data for a category."""
    sql = f"""
    SELECT year, value, count_ims
    FROM ttb_spirits_stats
    WHERE month IS NULL
    AND statistical_detail = '{category}'
    AND statistical_group LIKE '1-Distilled Spirits Production%'
    ORDER BY year DESC
    LIMIT {years}
    """
    results = run_d1_query(sql)
    return list(reversed(results)) if results else []


def create_yoy_chart(year, month, data):
    """Create year-over-year comparison bar chart."""
    # Organize data
    categories = {}
    for row in data:
        cat = row['statistical_detail']
        yr = row['year']
        if cat not in categories:
            categories[cat] = {'name': cat.split('-', 1)[1] if '-' in cat else cat}
        categories[cat][yr] = {'value': row['value'], 'producers': row['count_ims']}

    # Filter to categories with both years
    valid_cats = [c for c in categories if year in categories[c] and year-1 in categories[c]]
    if not valid_cats:
        return None, {}

    # Sort by current year value
    valid_cats.sort(key=lambda c: categories[c][year]['value'], reverse=True)

    # Prepare plot data
    names = [categories[c]['name'] for c in valid_cats]
    prior_vals = [categories[c][year-1]['value'] for c in valid_cats]
    current_vals = [categories[c][year]['value'] for c in valid_cats]
    colors = [get_category_color(c) for c in valid_cats]

    # Calculate changes
    changes = {}
    for c in valid_cats:
        pct = (categories[c][year]['value'] - categories[c][year-1]['value']) / categories[c][year-1]['value'] * 100
        prod_change = categories[c][year]['producers'] - categories[c][year-1]['producers']
        changes[categories[c]['name']] = {
            'pct': pct,
            'current': categories[c][year]['value'],
            'prior': categories[c][year-1]['value'],
            'producers_current': categories[c][year]['producers'],
            'producers_prior': categories[c][year-1]['producers'],
            'producers_change': prod_change
        }

    # Create figure
    fig, ax = plt.subplots(figsize=(10, 5.5))

    x = np.arange(len(names))
    width = 0.38

    # Bars
    bars_prior = ax.bar(x - width/2, prior_vals, width, label=str(year-1),
                        color='#B0B0B0', edgecolor='white', linewidth=1)
    bars_current = ax.bar(x + width/2, current_vals, width, label=str(year),
                          color=colors, edgecolor='white', linewidth=1)

    # Labels and formatting
    ax.set_xticks(x)
    ax.set_xticklabels(names, fontweight='medium')
    ax.yaxis.set_major_formatter(ticker.FuncFormatter(lambda x, p: format_number(x)))
    ax.set_ylabel('Proof Gallons', fontweight='medium')

    # Add change annotations
    for i, (bar, name) in enumerate(zip(bars_current, names)):
        change = changes[name]['pct']
        color = COLORS['positive'] if change > 0 else COLORS['negative']
        sign = '+' if change > 0 else ''

        y_pos = bar.get_height() + max(current_vals) * 0.02
        ax.annotate(f'{sign}{change:.1f}%',
                   xy=(bar.get_x() + bar.get_width()/2, y_pos),
                   ha='center', va='bottom',
                   fontsize=11, fontweight='bold', color=color)

    # Title
    month_name = datetime(2000, month, 1).strftime('%B')
    ax.set_title(f'Beverage Spirits Production: {month_name} {year} vs {year-1}',
                 pad=20, fontsize=14)

    # Legend
    ax.legend(loc='upper right', framealpha=0.9)

    # Subtle grid
    ax.yaxis.grid(True, alpha=0.3)
    ax.set_axisbelow(True)

    # Source annotation
    fig.text(0.99, 0.01, 'Source: TTB Distilled Spirits Statistics',
             ha='right', va='bottom', fontsize=8, color='#888888', style='italic')

    plt.tight_layout()

    # Save to buffer
    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=200, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    buf.seek(0)
    plt.close()

    return buf, changes


def create_trend_chart(category, data):
    """Create multi-year trend chart."""
    if not data:
        return None

    years = [r['year'] for r in data]
    values = [r['value'] for r in data]
    producers = [r['count_ims'] for r in data]

    # Get category display name and color
    display_name = category.split('-', 1)[1] if '-' in category else category
    color = get_category_color(category)

    # Create figure with dual y-axis
    fig, ax1 = plt.subplots(figsize=(10, 5.5))
    ax2 = ax1.twinx()

    # Production area and line
    ax1.fill_between(years, values, alpha=0.15, color=color)
    line1 = ax1.plot(years, values, color=color, linewidth=2.5, marker='o',
                     markersize=8, markerfacecolor='white', markeredgewidth=2,
                     label='Production Volume')

    # Producer count line
    line2 = ax2.plot(years, producers, color=COLORS['secondary'], linewidth=2,
                     marker='s', markersize=6, linestyle='--',
                     label='Active Producers')

    # Format axes
    ax1.yaxis.set_major_formatter(ticker.FuncFormatter(lambda x, p: format_number(x)))
    ax1.set_ylabel('Production (Proof Gallons)', color=color, fontweight='medium')
    ax1.tick_params(axis='y', labelcolor=color)

    ax2.set_ylabel('Number of Producers', color=COLORS['secondary'], fontweight='medium')
    ax2.tick_params(axis='y', labelcolor=COLORS['secondary'])
    ax2.spines['right'].set_visible(True)
    ax2.spines['right'].set_color(COLORS['secondary'])
    ax2.spines['right'].set_alpha(0.5)

    ax1.set_xlabel('Year', fontweight='medium')
    ax1.set_xticks(years)

    # Add value labels
    for x, y in zip(years, values):
        ax1.annotate(format_number(y, 0), (x, y),
                    textcoords="offset points", xytext=(0, 12),
                    ha='center', fontsize=9, color=color, fontweight='medium')

    # Title
    ax1.set_title(f'{display_name} Production Trend ({years[0]}-{years[-1]})',
                  pad=15, fontsize=14)

    # Combined legend
    lines1, labels1 = ax1.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(lines1 + lines2, labels1 + labels2, loc='upper left', framealpha=0.9)

    # Grid
    ax1.yaxis.grid(True, alpha=0.3)
    ax1.set_axisbelow(True)

    # Source
    fig.text(0.99, 0.01, 'Source: TTB Distilled Spirits Statistics',
             ha='right', va='bottom', fontsize=8, color='#888888', style='italic')

    plt.tight_layout()

    # Save to buffer
    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=200, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    buf.seek(0)
    plt.close()

    return buf


def create_production_breakdown_chart(year, month, data):
    """Create production breakdown showing beverage vs industrial."""
    if not data:
        return None

    # Separate beverage and industrial
    beverage = []
    industrial = None

    for row in data:
        cat = row['statistical_detail']
        if 'Neutral Spirits' in cat or 'Alcohol' in cat:
            industrial = row
        else:
            beverage.append(row)

    # Sort beverage by value
    beverage.sort(key=lambda x: x['value'], reverse=True)

    # Create figure with two subplots
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5),
                                    gridspec_kw={'width_ratios': [1.2, 1]})

    # Left: All categories horizontal bar
    all_cats = beverage + ([industrial] if industrial else [])
    names = []
    values = []
    colors = []

    for row in all_cats:
        cat = row['statistical_detail']
        name = cat.split('-', 1)[1] if '-' in cat else cat
        # Shorten long names
        if 'Alcohol, Neutral' in name:
            name = 'Neutral Spirits'
        names.append(name)
        values.append(row['value'])
        colors.append(get_category_color(cat))

    y_pos = np.arange(len(names))
    bars = ax1.barh(y_pos, values, color=colors, edgecolor='white', height=0.7)

    ax1.set_yticks(y_pos)
    ax1.set_yticklabels(names, fontweight='medium')
    ax1.invert_yaxis()
    ax1.xaxis.set_major_formatter(ticker.FuncFormatter(lambda x, p: format_number(x)))
    ax1.set_xlabel('Proof Gallons', fontweight='medium')
    ax1.set_title('All Distilled Spirits Production', fontsize=12, fontweight='bold')

    # Add value labels
    for bar, val in zip(bars, values):
        label = format_number(val)
        ax1.text(bar.get_width() + max(values) * 0.02,
                bar.get_y() + bar.get_height()/2,
                label, va='center', fontsize=9, fontweight='medium')

    # Right: Beverage-only with producer counts
    if beverage:
        bev_names = []
        bev_values = []
        bev_producers = []
        bev_colors = []

        for row in beverage:
            cat = row['statistical_detail']
            name = cat.split('-', 1)[1] if '-' in cat else cat
            bev_names.append(name)
            bev_values.append(row['value'])
            bev_producers.append(row['count_ims'])
            bev_colors.append(get_category_color(cat))

        y_pos2 = np.arange(len(bev_names))
        bars2 = ax2.barh(y_pos2, bev_values, color=bev_colors, edgecolor='white', height=0.7)

        ax2.set_yticks(y_pos2)
        ax2.set_yticklabels(bev_names, fontweight='medium')
        ax2.invert_yaxis()
        ax2.xaxis.set_major_formatter(ticker.FuncFormatter(lambda x, p: format_number(x)))
        ax2.set_xlabel('Proof Gallons', fontweight='medium')
        ax2.set_title('Beverage Spirits Detail', fontsize=12, fontweight='bold')

        # Add value + producer labels
        for bar, val, prod in zip(bars2, bev_values, bev_producers):
            label = f'{format_number(val)} ({prod:,} producers)'
            ax2.text(bar.get_width() + max(bev_values) * 0.02,
                    bar.get_y() + bar.get_height()/2,
                    label, va='center', fontsize=9)

    # Main title
    month_name = datetime(2000, month, 1).strftime('%B')
    fig.suptitle(f'Distilled Spirits Production: {month_name} {year}',
                 fontsize=14, fontweight='bold', y=1.02)

    # Source
    fig.text(0.99, 0.01, 'Source: TTB Distilled Spirits Statistics',
             ha='right', va='bottom', fontsize=8, color='#888888', style='italic')

    plt.tight_layout()

    # Save to buffer
    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=200, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    buf.seek(0)
    plt.close()

    return buf


def generate_commentary(year, month, monthly_data, yoy_changes, trend_data):
    """Generate analysis commentary based on the data."""
    month_name = datetime(2000, month, 1).strftime('%B')

    commentary = {
        'headline': '',
        'summary': '',
        'whisky_analysis': '',
        'other_categories': '',
        'outlook': ''
    }

    # Find key metrics
    whisky_change = yoy_changes.get('Whisky', {})
    brandy_change = yoy_changes.get('Brandy', {})
    rgv_change = yoy_changes.get('Rum, Gin, & Vodka', {})

    # Headline
    if whisky_change:
        whisky_pct = whisky_change.get('pct', 0)
        if whisky_pct < -20:
            commentary['headline'] = f'American Whisky Production Plunges {abs(whisky_pct):.0f}% in {month_name} {year}'
        elif whisky_pct < -10:
            commentary['headline'] = f'Whisky Production Down {abs(whisky_pct):.0f}% as Industry Contracts'
        elif whisky_pct < 0:
            commentary['headline'] = f'Whisky Production Slips {abs(whisky_pct):.1f}% in {month_name}'
        elif whisky_pct > 10:
            commentary['headline'] = f'Whisky Production Surges {whisky_pct:.0f}% in {month_name}'
        else:
            commentary['headline'] = f'{month_name} {year} Spirits Production Report'

    # Summary paragraph
    total_beverage = sum(d['value'] for d in monthly_data
                        if 'Neutral' not in d['statistical_detail'] and 'Alcohol' not in d['statistical_detail'])

    commentary['summary'] = (
        f"American distillers produced {format_number(total_beverage)} proof gallons of beverage spirits in {month_name} {year}. "
    )

    if whisky_change:
        commentary['summary'] += (
            f"Whisky accounted for {format_number(whisky_change.get('current', 0))} proof gallons, "
            f"{'down' if whisky_change['pct'] < 0 else 'up'} {abs(whisky_change['pct']):.1f}% from {month_name} {year-1}. "
        )

        prod_change = whisky_change.get('producers_change', 0)
        if prod_change != 0:
            commentary['summary'] += (
                f"The number of active whisky producers {'fell' if prod_change < 0 else 'rose'} to "
                f"{whisky_change.get('producers_current', 0):,}, {'down' if prod_change < 0 else 'up'} "
                f"{abs(prod_change)} from the prior year."
            )

    # Whisky analysis
    if whisky_change and trend_data.get('whisky'):
        whisky_trend = trend_data['whisky']
        peak_year = max(whisky_trend, key=lambda x: x['value'])

        commentary['whisky_analysis'] = (
            f"The whisky production decline extends a trend that began after the {peak_year['year']} peak of "
            f"{format_number(peak_year['value'])} proof gallons. "
        )

        if whisky_change['pct'] < -15:
            commentary['whisky_analysis'] += (
                "The sharp pullback reflects elevated inventory levels across the bourbon sector. "
                "With aging whisky tying up capital and warehouse capacity, producers have scaled back output "
                "while working through existing stock."
            )
        elif whisky_change['pct'] < 0:
            commentary['whisky_analysis'] += (
                "Industry observers attribute the decline to inventory normalization following several years "
                "of aggressive production expansion."
            )

    # Other categories
    other_parts = []
    if brandy_change:
        direction = 'fell' if brandy_change['pct'] < 0 else 'rose'
        other_parts.append(
            f"Brandy production {direction} {abs(brandy_change['pct']):.0f}% to {format_number(brandy_change.get('current', 0))} "
            f"proof gallons from {brandy_change.get('producers_current', 0)} producers"
        )

    if rgv_change:
        direction = 'declined' if rgv_change['pct'] < 0 else 'increased'
        other_parts.append(
            f"The combined rum, gin, and vodka category {direction} {abs(rgv_change['pct']):.0f}% to "
            f"{format_number(rgv_change.get('current', 0))} proof gallons"
        )

    if other_parts:
        commentary['other_categories'] = '. '.join(other_parts) + '.'

    # Outlook
    if whisky_change and whisky_change['pct'] < -10:
        commentary['outlook'] = (
            "The production pullback suggests the industry is entering a consolidation phase after a decade of expansion. "
            "Whether this adjustment lasts one year or several will depend on how quickly inventory levels normalize "
            "and whether consumer demand growth resumes."
        )

    return commentary


def create_word_document(year, month, charts, commentary, data_summary):
    """Create Word document with charts and commentary."""
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading(commentary['headline'], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle with date
    month_name = datetime(2000, month, 1).strftime('%B')
    subtitle = doc.add_paragraph(f'TTB Distilled Spirits Statistics | {month_name} {year}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    subtitle.runs[0].font.size = Pt(12)

    doc.add_paragraph()

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(commentary['summary'])

    # Production Overview Chart
    if 'production' in charts and charts['production']:
        doc.add_heading('Production Overview', level=1)
        doc.add_picture(charts['production'], width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Year-over-Year Comparison
    if 'yoy' in charts and charts['yoy']:
        doc.add_heading('Year-over-Year Comparison', level=1)
        doc.add_picture(charts['yoy'], width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add YoY data table
        if data_summary.get('yoy_changes'):
            doc.add_paragraph()
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'

            # Header row
            headers = ['Category', f'{year-1}', f'{year}', 'Change', 'Producers']
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True

            # Data rows
            for cat, vals in data_summary['yoy_changes'].items():
                row = table.add_row()
                row.cells[0].text = cat
                row.cells[1].text = format_number(vals['prior'])
                row.cells[2].text = format_number(vals['current'])
                sign = '+' if vals['pct'] > 0 else ''
                row.cells[3].text = f"{sign}{vals['pct']:.1f}%"
                prod_sign = '+' if vals['producers_change'] > 0 else ''
                row.cells[4].text = f"{vals['producers_current']:,} ({prod_sign}{vals['producers_change']})"

    # Whisky Analysis
    if commentary.get('whisky_analysis'):
        doc.add_heading('Whisky Analysis', level=1)
        doc.add_paragraph(commentary['whisky_analysis'])

        # Add whisky trend chart
        if 'whisky_trend' in charts and charts['whisky_trend']:
            doc.add_picture(charts['whisky_trend'], width=Inches(6.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Other Categories
    if commentary.get('other_categories'):
        doc.add_heading('Other Categories', level=1)
        doc.add_paragraph(commentary['other_categories'])

        # Add brandy trend if significant decline
        if 'brandy_trend' in charts and charts['brandy_trend']:
            doc.add_picture(charts['brandy_trend'], width=Inches(6.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Outlook
    if commentary.get('outlook'):
        doc.add_heading('Outlook', level=1)
        doc.add_paragraph(commentary['outlook'])

    # Data Source
    doc.add_paragraph()
    source = doc.add_paragraph('Source: TTB Distilled Spirits Statistics')
    source.runs[0].font.italic = True
    source.runs[0].font.size = Pt(9)
    source.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    return doc


def main():
    parser = argparse.ArgumentParser(description='Generate spirits statistics report')
    parser.add_argument('--year', type=int, help='Year')
    parser.add_argument('--month', type=int, help='Month (1-12)')
    args = parser.parse_args()

    # Get latest data if not specified
    if not args.year or not args.month:
        sql = "SELECT MAX(year) as y, MAX(month) as m FROM ttb_spirits_stats WHERE month IS NOT NULL"
        results = run_d1_query(sql)
        if results:
            args.year = args.year or results[0]['y']
            args.month = args.month or results[0]['m']
        else:
            print("No data found. Run import first.")
            return

    print(f"\nGenerating report for {args.year}-{args.month:02d}...")

    # Fetch data
    print("  Fetching monthly data...")
    monthly_data = get_monthly_data(args.year, args.month)
    if not monthly_data:
        print("No monthly data found.")
        return

    print("  Fetching YoY comparison...")
    yoy_data = get_yoy_data(args.year, args.month)

    print("  Fetching trend data...")
    trend_data = {
        'whisky': get_yearly_trend('1-Whisky'),
        'brandy': get_yearly_trend('2-Brandy'),
        'rgv': get_yearly_trend('3-Rum, Gin, & Vodka')
    }

    # Generate charts
    print("  Creating charts...")
    charts = {}

    # Production breakdown
    charts['production'] = create_production_breakdown_chart(args.year, args.month, monthly_data)

    # YoY comparison
    yoy_chart, yoy_changes = create_yoy_chart(args.year, args.month, yoy_data)
    charts['yoy'] = yoy_chart

    # Trend charts
    if trend_data['whisky']:
        charts['whisky_trend'] = create_trend_chart('1-Whisky', trend_data['whisky'])
    if trend_data['brandy']:
        charts['brandy_trend'] = create_trend_chart('2-Brandy', trend_data['brandy'])

    # Generate commentary
    print("  Generating commentary...")
    commentary = generate_commentary(args.year, args.month, monthly_data, yoy_changes, trend_data)

    # Create data summary
    data_summary = {
        'yoy_changes': yoy_changes,
        'monthly_data': monthly_data
    }

    # Create Word document
    print("  Creating Word document...")
    doc = create_word_document(args.year, args.month, charts, commentary, data_summary)

    # Save
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    month_name = datetime(2000, args.month, 1).strftime('%B').lower()
    output_path = os.path.join(OUTPUT_DIR, f'spirits-report-{args.year}-{args.month:02d}-{month_name}.docx')
    doc.save(output_path)

    print(f"\nReport saved to: {output_path}")

    # Also save charts as standalone files
    chart_dir = os.path.join(OUTPUT_DIR, 'charts')
    os.makedirs(chart_dir, exist_ok=True)

    for name, buf in charts.items():
        if buf:
            buf.seek(0)
            chart_path = os.path.join(chart_dir, f'{name}-{args.year}-{args.month:02d}.png')
            with open(chart_path, 'wb') as f:
                f.write(buf.read())
            print(f"  Chart saved: {chart_path}")


if __name__ == '__main__':
    main()
